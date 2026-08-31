import asyncio
from concurrent.futures import ThreadPoolExecutor
from dataclasses import replace
import io
import json
import os
from pathlib import Path
import signal
import sys
import time
import zipfile

import psutil
import pytest

from packages.knowledge.errors import KnowledgeValidationError, ParseLimitExceeded, ParseTimeout, ParseProtocolError, ParseUnavailable
from packages.knowledge.ingestion.isolated import IsolatedDocumentParser
from packages.knowledge.ingestion.limits import ParseLimits
from packages.knowledge.ingestion.parsers import DocumentParser
from packages.knowledge.ingestion.chunker import StructureAwareChunker
from packages.knowledge.models import ParsedBlock


def fixture_parser(monkeypatch, mode, limits=None, *args):
    parser = IsolatedDocumentParser(limits)
    monkeypatch.setattr(parser, '_command', lambda: [sys.executable, '-I', '-B',
        str(Path(__file__).with_name('parser_fixture_worker.py')), mode, *map(str, args)])
    processes = []
    original = parser._spawn
    async def spawn(directory):
        process = await original(directory)
        processes.append(process)
        return process
    monkeypatch.setattr(parser, '_spawn', spawn)
    return parser, processes


def pdf(pages=1):
    from pypdf import PdfWriter
    from pypdf.generic import DictionaryObject, NameObject, DecodedStreamObject
    writer = PdfWriter()
    font = DictionaryObject({NameObject('/Type'): NameObject('/Font'), NameObject('/Subtype'): NameObject('/Type1'),
        NameObject('/BaseFont'): NameObject('/Helvetica')})
    for _ in range(pages):
        page = writer.add_blank_page(width=200, height=200)
        page[NameObject('/Resources')] = DictionaryObject({NameObject('/Font'): DictionaryObject({NameObject('/F1'): writer._add_object(font)})})
        stream = DecodedStreamObject()
        stream.set_data(b'BT /F1 12 Tf 10 50 Td (Verified document text.) Tj ET')
        page[NameObject('/Contents')] = writer._add_object(stream)
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


@pytest.mark.parametrize('content,mime,name,locator', [
    ('# 标题\n\n正文。'.encode(), 'text/markdown', 'a.md', {'section': '标题'}),
    (b'First paragraph\n\nSecond paragraph', 'text/plain', 'a.txt', {'paragraph': 1}),
    (b'a,b\nc,d', 'text/csv', 'a.csv', {'row': 1}),
    (b'{"key": "value"}', 'application/json', 'a.json', {'section': 'root'}),
    (b'<h1>Heading</h1><p>body</p>', 'text/html', 'a.html', {'section': 'Heading'}),
])
def test_real_subprocess_formats_and_citations(content, mime, name, locator):
    chunks = asyncio.run(IsolatedDocumentParser().parse(content, mime, name))
    assert chunks and chunks[0].locator == locator
    assert all(len(chunk.text) <= 2200 for chunk in chunks)


def test_actual_docx_and_pdf_child_parsing():
    from docx import Document
    document = Document()
    document.add_heading('Document heading', 1)
    document.add_paragraph('Verified document text.')
    stream = io.BytesIO()
    document.save(stream)
    parser = IsolatedDocumentParser()
    docx_chunks = asyncio.run(parser.parse(stream.getvalue(), 'application/octet-stream', 'a.docx'))
    assert docx_chunks[1].text == 'Verified document text.'
    assert docx_chunks[1].locator['section'] == 'Document heading'
    pdf_chunks = asyncio.run(parser.parse(pdf(), 'application/pdf', 'a.pdf'))
    assert pdf_chunks[0].text == 'Verified document text.' and pdf_chunks[0].locator == {'page': 1}


@pytest.mark.parametrize('field,value', [
    ('timeout_seconds', float('nan')), ('timeout_seconds', 0), ('memory_bytes', -1), ('max_chunks', True),
    ('max_pages', 1.5), ('overlap_characters', 2200), ('chunk_characters', 0), ('max_concurrent', 100),
])
def test_bad_limits_rejected(field, value):
    with pytest.raises(ValueError):
        ParseLimits(**{field: value})


@pytest.mark.parametrize('field,value,content,mime,name', [
    ('max_input_bytes', 3, b'1234', 'text/plain', 'a.txt'),
    ('max_text_characters', 3, b'1234', 'text/plain', 'a.txt'),
    ('max_blocks', 1, b'a\n\nb', 'text/plain', 'a.txt'),
    ('max_chunks', 1, b'a' * 3000, 'text/plain', 'a.txt'),
    ('max_output_bytes', 50, b'a', 'text/plain', 'a.txt'),
])
def test_limits_fail_without_silent_truncation(field, value, content, mime, name):
    parser = IsolatedDocumentParser(replace(ParseLimits(), **{field: value}))
    with pytest.raises(ParseLimitExceeded):
        asyncio.run(parser.parse(content, mime, name))


def test_pdf_page_and_expanded_limits():
    with pytest.raises(ParseLimitExceeded):
        asyncio.run(IsolatedDocumentParser(ParseLimits(max_pages=1)).parse(pdf(2), 'application/pdf', 'a.pdf'))
    with pytest.raises(ParseLimitExceeded):
        asyncio.run(IsolatedDocumentParser(ParseLimits(max_expanded_bytes=10)).parse(pdf(), 'application/pdf', 'a.pdf'))


def archive(items, compression=zipfile.ZIP_DEFLATED):
    output = io.BytesIO()
    with zipfile.ZipFile(output, 'w', compression=compression) as container:
        for name, value in items:
            container.writestr(name, value)
    return output.getvalue()


def test_zip_expansion_ratio_entries_and_ambiguous_names():
    bomb = archive([('word/document.xml', b'0' * 1_000_000)])
    for limits in (ParseLimits(max_expanded_bytes=1000), ParseLimits(max_compression_ratio=2)):
        with pytest.raises(ParseLimitExceeded):
            asyncio.run(IsolatedDocumentParser(limits).parse(bomb, 'application/octet-stream', 'a.docx'))
    entries = archive([('first', b'a'), ('second', b'b')])
    with pytest.raises(ParseLimitExceeded):
        asyncio.run(IsolatedDocumentParser(ParseLimits(max_archive_entries=1)).parse(entries, 'application/octet-stream', 'a.docx'))
    with pytest.warns(UserWarning, match='Duplicate'):
        duplicate = archive([('same', b'a'), ('same', b'b')])
    with pytest.raises(KnowledgeValidationError):
        asyncio.run(IsolatedDocumentParser().parse(duplicate, 'application/octet-stream', 'a.docx'))


@pytest.mark.parametrize('mode,error', [
    ('hang', ParseTimeout), ('blocked-input', ParseTimeout), ('stdout', ParseLimitExceeded),
    ('stderr', ParseLimitExceeded), ('crash', ParseLimitExceeded), ('protocol', ParseProtocolError),
    ('orphan', ParseTimeout),
])
def test_child_failure_is_bounded_reaped_and_followed_by_success(monkeypatch, mode, error):
    parser, processes = fixture_parser(monkeypatch, mode, ParseLimits(timeout_seconds=1, max_output_bytes=8192))
    async def scenario():
        start = time.monotonic()
        with pytest.raises(error) as caught:
            await parser.parse(b'a' * 256_000, 'text/plain', 'a.txt')
        assert 'private-document-sentinel' not in str(caught.value)
        assert time.monotonic() - start < 5
        assert processes and all(item.returncode is not None for item in processes)
        result = await IsolatedDocumentParser().parse(b'next document', 'text/plain', 'b.txt')
        assert result[0].text == 'next document'
    asyncio.run(scenario())


def test_resident_memory_watchdog_kills_real_child(monkeypatch):
    parser, processes = fixture_parser(monkeypatch, 'memory', ParseLimits(memory_bytes=48 * 1024**2, timeout_seconds=5))
    with pytest.raises(ParseLimitExceeded):
        asyncio.run(parser.parse(b'input', 'text/plain', 'a.txt'))
    assert processes[0].returncode is not None


def test_cpu_hard_limit_applies_to_real_child(monkeypatch):
    parser, processes = fixture_parser(monkeypatch, 'cpu', ParseLimits(cpu_seconds=1, timeout_seconds=5))
    with pytest.raises((ParseTimeout, ParseLimitExceeded)):
        asyncio.run(parser.parse(b'input', 'text/plain', 'a.txt'))
    assert processes[0].returncode in {-signal.SIGXCPU, -signal.SIGKILL}


def test_cancellation_reaps_process_group_and_does_not_block_event_loop(monkeypatch):
    parser, processes = fixture_parser(monkeypatch, 'child')
    async def scenario():
        task = asyncio.create_task(parser.parse(b'input', 'text/plain', 'a.txt'))
        children = []
        try:
            async with asyncio.timeout(5):
                while not children:
                    await asyncio.sleep(0.01)
                    if processes:
                        children = psutil.Process(processes[0].pid).children()
            task.cancel()
            with pytest.raises(asyncio.CancelledError):
                await task
            assert processes[0].returncode is not None
            assert all(not child.is_running() or child.status() == psutil.STATUS_ZOMBIE for child in children)
        finally:
            task.cancel()
            await asyncio.gather(task, return_exceptions=True)
    asyncio.run(scenario())


def test_cancellation_during_spawn_still_reaps_child(monkeypatch):
    parser = IsolatedDocumentParser()
    original = parser._spawn
    processes = []
    async def scenario():
        spawned = asyncio.Event()
        async def delayed(directory):
            process = await original(directory)
            processes.append(process)
            spawned.set()
            await asyncio.sleep(.1)
            return process
        monkeypatch.setattr(parser, '_spawn', delayed)
        task = asyncio.create_task(parser.parse(b'input', 'text/plain', 'a.txt'))
        await spawned.wait()
        task.cancel()
        await asyncio.sleep(.02)
        task.cancel()  # repeated cancellation must not abandon cleanup
        with pytest.raises(asyncio.CancelledError):
            await task
        assert processes[0].returncode is not None
    asyncio.run(scenario())


def test_environment_fds_and_cwd_are_not_inherited(monkeypatch, tmp_path):
    monkeypatch.setenv('OPENAI_API_KEY', 'secret-not-for-parser')
    monkeypatch.setenv('DATABASE_URL', 'secret-db-not-for-parser')
    path = tmp_path / 'parent-file'
    path.write_text('private parent data')
    with path.open('rb') as stream:
        os.set_inheritable(stream.fileno(), True)
        parser, _ = fixture_parser(monkeypatch, 'inspect', None, stream.fileno())
        data = json.loads(asyncio.run(parser.parse(b'input', 'text/plain', 'a.txt'))[0].text)
    assert set(data['environment']) <= {'PATH', 'LANG', 'LC_ALL', 'OMP_NUM_THREADS', 'OPENBLAS_NUM_THREADS'}
    assert data['isolated'] and data['no_bytecode'] and not data['inherited_fd']
    assert not Path(data['cwd']).exists()


def test_reused_parser_across_event_loops_and_concurrent_threads():
    parser = IsolatedDocumentParser(ParseLimits(max_concurrent=1))
    with ThreadPoolExecutor(max_workers=3) as pool:
        results = list(pool.map(lambda index: asyncio.run(parser.parse(f'document {index}'.encode(), 'text/plain', 'a.txt')), range(3)))
    assert [chunks[0].text for chunks in results] == ['document 0', 'document 1', 'document 2']


def test_chunk_boundary_and_zero_overlap():
    chunker = StructureAwareChunker(10, 3)
    chunks = chunker.chunk([ParsedBlock(text='12345678. abcdefghi. abcdefghijklmnop')])
    assert all(len(chunk.text) <= 10 for chunk in chunks)
    assert [chunk.text for chunk in StructureAwareChunker(5, 0).chunk([ParsedBlock(text='abcdefghij')])] == ['abcde', 'fghij']


def test_production_refuses_missing_linux_hard_memory_boundary(monkeypatch):
    monkeypatch.setenv('DEEPAGENT_ENVIRONMENT', 'production')
    monkeypatch.setattr(sys, 'platform', 'darwin')
    with pytest.raises(ParseUnavailable, match='requires Linux'):
        asyncio.run(IsolatedDocumentParser().parse(b'input', 'text/plain', 'a.txt'))
