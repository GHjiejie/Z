from __future__ import annotations

import csv
import io
import json
import re
import zipfile
from html.parser import HTMLParser
from pathlib import PurePosixPath
from typing import List

from packages.knowledge.errors import KnowledgeValidationError, ParseLimitExceeded
from packages.knowledge.models import ParsedBlock
from .limits import ParseLimits


class _Blocks(list):
    def __init__(self, limits):
        super().__init__()
        self.limits, self.characters = limits, 0

    def append(self, block):
        self.characters += len(block.text)
        if len(self) >= self.limits.max_blocks or self.characters > self.limits.max_text_characters:
            raise ParseLimitExceeded("Document exceeds extracted text or block limits")
        if any(isinstance(value, str) and len(value) > 512 for value in block.locator.values()):
            raise ParseLimitExceeded("Document location metadata exceeds the limit")
        super().append(block)


class _TextHTMLParser(HTMLParser):
    def __init__(self, limits):
        super().__init__()
        self.parts: List[str] = []
        self.section: str | None = None
        self.blocks = _Blocks(limits)

    def handle_starttag(self, tag, attrs):
        if tag in {"h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "tr", "br"}:
            self._flush()

    def handle_endtag(self, tag):
        if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            text = " ".join(self.parts).strip()
            if text:
                self.section = text
            self._flush()
        elif tag in {"p", "li", "tr"}:
            self._flush()

    def handle_data(self, data):
        if data.strip():
            self.parts.append(data.strip())

    def close(self):
        super().close()
        self._flush()

    def _flush(self):
        text = " ".join(self.parts).strip()
        if text:
            locator = {"section": self.section} if self.section else {}
            self.blocks.append(ParsedBlock(text=text, locator=locator))
        self.parts = []


class DocumentParser:
    version = "structure-parser-2.0"

    def __init__(self, limits: ParseLimits | None = None):
        self.limits = limits or ParseLimits()

    def parse(self, content: bytes, content_type: str, filename: str) -> List[ParsedBlock]:
        if len(content) > self.limits.max_input_bytes:
            raise ParseLimitExceeded("Document exceeds input size limit")
        suffix = PurePosixPath(filename).suffix.lower()
        normalized_type = content_type.split(";", 1)[0].strip().lower()
        if suffix == ".pdf" or normalized_type == "application/pdf":
            return self._parse_pdf(content)
        if suffix == ".docx" or normalized_type.endswith("wordprocessingml.document"):
            return self._parse_docx(content)
        if suffix in {".html", ".htm"} or normalized_type == "text/html":
            parser = _TextHTMLParser(self.limits)
            parser.feed(self._decode(content))
            parser.close()
            return self._require_content(parser.blocks)
        if suffix == ".json" or normalized_type == "application/json":
            try:
                value = json.loads(self._decode(content))
                text = json.dumps(value, ensure_ascii=False, indent=2)
            except json.JSONDecodeError as exc:
                raise KnowledgeValidationError(f"Invalid JSON document: {exc}") from exc
            blocks = _Blocks(self.limits)
            blocks.append(ParsedBlock(text=text, locator={"section": "root"}))
            return blocks
        if suffix == ".csv" or normalized_type == "text/csv":
            return self._parse_csv(content)
        if suffix in {".md", ".markdown"} or normalized_type in {
            "text/markdown",
            "text/x-markdown",
        }:
            return self._parse_markdown(self._decode(content))
        if suffix in {".txt", ""} or normalized_type.startswith("text/"):
            return self._parse_plain_text(self._decode(content))
        raise KnowledgeValidationError(
            f"Unsupported document type: {content_type or suffix or 'unknown'}"
        )

    def _parse_pdf(self, content: bytes) -> List[ParsedBlock]:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise KnowledgeValidationError("PDF parsing requires the pypdf dependency") from exc
        try:
            reader = PdfReader(io.BytesIO(content))
            if len(reader.pages) > self.limits.max_pages:
                raise ParseLimitExceeded("PDF exceeds page count limit")
            blocks = _Blocks(self.limits)
            expanded = 0
            for index, page in enumerate(reader.pages, start=1):
                stream = page.get_contents()
                if stream is not None:
                    expanded += len(stream.get_data())
                if expanded > self.limits.max_expanded_bytes:
                    raise ParseLimitExceeded("PDF exceeds expanded content limit")
                text = (page.extract_text() or "").strip()
                if text:
                    blocks.append(ParsedBlock(text=text, locator={"page": index}))
            return self._require_content(blocks)
        except (ParseLimitExceeded, MemoryError):
            raise
        except Exception as exc:
            raise KnowledgeValidationError("Unable to parse PDF document") from exc

    def _parse_docx(self, content: bytes) -> List[ParsedBlock]:
        try:
            from docx import Document
        except ImportError as exc:
            raise KnowledgeValidationError("DOCX parsing requires the python-docx dependency") from exc
        try:
            self._validate_docx_archive(content)
            document = Document(io.BytesIO(content))
            blocks = _Blocks(self.limits)
            section: str | None = None
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                if paragraph.style and paragraph.style.name.lower().startswith("heading"):
                    section = text
                blocks.append(
                    ParsedBlock(text=text, locator={"section": section} if section else {})
                )
            for table_index, table in enumerate(document.tables, start=1):
                rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
                text = "\n".join(row for row in rows if row.strip(" |"))
                if text:
                    blocks.append(
                        ParsedBlock(text=text, locator={"table": table_index, "section": section})
                    )
            return self._require_content(blocks)
        except (ParseLimitExceeded, MemoryError):
            raise
        except Exception as exc:
            raise KnowledgeValidationError("Unable to parse DOCX document") from exc

    def _validate_docx_archive(self, content):
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            entries = archive.infolist()
            if len(entries) > self.limits.max_archive_entries:
                raise ParseLimitExceeded("DOCX exceeds archive entry limit")
            expanded = 0
            seen = set()
            for entry in entries:
                path = PurePosixPath(entry.filename)
                if (entry.filename in seen or path.is_absolute() or '..' in path.parts or '\\' in entry.filename
                        or entry.flag_bits & 1 or entry.compress_type not in (zipfile.ZIP_STORED, zipfile.ZIP_DEFLATED)):
                    raise KnowledgeValidationError("Unsupported or ambiguous DOCX archive")
                seen.add(entry.filename)
                expanded += entry.file_size
                if expanded > self.limits.max_expanded_bytes:
                    raise ParseLimitExceeded("DOCX exceeds expanded size limit")
                if entry.file_size > max(1, entry.compress_size) * self.limits.max_compression_ratio:
                    raise ParseLimitExceeded("DOCX exceeds compression ratio limit")
            actual = 0
            for entry in entries:
                with archive.open(entry) as stream:
                    while data := stream.read(min(65536, self.limits.max_expanded_bytes - actual + 1)):
                        actual += len(data)
                        if actual > self.limits.max_expanded_bytes:
                            raise ParseLimitExceeded("DOCX exceeds expanded size limit")

    def _parse_markdown(self, text: str) -> List[ParsedBlock]:
        blocks = _Blocks(self.limits)
        section: str | None = None
        buffer: List[str] = []

        def flush():
            if buffer:
                value = "\n".join(buffer).strip()
                if value:
                    blocks.append(
                        ParsedBlock(text=value, locator={"section": section} if section else {})
                    )
                buffer.clear()

        for line in text.splitlines():
            heading = re.match(r"^#{1,6}\s+(.+)$", line.strip())
            if heading:
                flush()
                section = heading.group(1).strip()
                buffer.append(line.strip())
            elif line.strip():
                buffer.append(line.rstrip())
            else:
                flush()
        flush()
        return self._require_content(blocks)

    def _parse_plain_text(self, text: str) -> List[ParsedBlock]:
        blocks = _Blocks(self.limits)
        for index, part in enumerate(re.split(r"\n\s*\n", text), start=1):
            if part.strip():
                blocks.append(ParsedBlock(text=part.strip(), locator={"paragraph": index}))
        return self._require_content(blocks)

    def _parse_csv(self, content: bytes) -> List[ParsedBlock]:
        csv.field_size_limit(self.limits.max_text_characters)
        reader = csv.reader(io.StringIO(self._decode(content)))
        blocks = _Blocks(self.limits)
        for index, row in enumerate(reader, start=1):
            text = " | ".join(cell.strip() for cell in row)
            if text.strip(" |"):
                blocks.append(ParsedBlock(text=text, locator={"row": index}))
        return self._require_content(blocks)

    def _decode(self, content: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                text = content.decode(encoding)
                if len(text) > self.limits.max_text_characters:
                    raise ParseLimitExceeded("Document exceeds text character limit")
                return text
            except UnicodeDecodeError:
                continue
        raise KnowledgeValidationError("Document text encoding is not supported")

    @staticmethod
    def _require_content(blocks: List[ParsedBlock]) -> List[ParsedBlock]:
        if not blocks:
            raise KnowledgeValidationError("Document contains no extractable text")
        return blocks
